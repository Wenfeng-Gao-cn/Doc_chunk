"""
readfile_2_str.py - 文件读取模块

支持读取以下文件格式并转换为字符串：
- Word文档 (.docx, .doc)
- Excel文档 (.xlsx, .xls)
- Markdown文档 (.md)
- 文本文档 (.txt)
- PDF文档 (.pdf)
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any


def read_txt_file(file_path: str) -> str:
    """
    读取文本文件
    
    Args:
        file_path (str): 文件路径
        
    Returns:
        str: 文件内容
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # 如果UTF-8编码失败，尝试其他编码
        encodings = ['gbk', 'gb2312', 'ascii', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        raise Exception(f"无法解码文件 {file_path}")


def read_markdown_file(file_path: str) -> str:
    """
    读取Markdown文件
    
    Args:
        file_path (str): 文件路径
        
    Returns:
        str: 文件内容
    """
    return read_txt_file(file_path)


def read_word_file(file_path: str) -> str:
    """
    读取Word文档
    
    Args:
        file_path (str): 文件路径
        
    Returns:
        str: 文件内容
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        content = []
        
        # 读取段落
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        
        # 读取表格
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                content.append('\t'.join(row_data))
        
        return '\n'.join(content)
    
    except ImportError:
        raise ImportError("需要安装python-docx库: pip install python-docx")
    except Exception as e:
        raise Exception(f"读取Word文档失败: {str(e)}")


def read_excel_file(file_path: str) -> str:
    """
    读取Excel文档
    
    Args:
        file_path (str): 文件路径
        
    Returns:
        str: 文件内容，包含所有工作表
    """
    try:
        import pandas as pd
        
        # 读取所有工作表
        excel_file = pd.ExcelFile(file_path)
        content = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            content.append(f"=== 工作表: {sheet_name} ===")
            content.append(df.to_string(index=False))
            content.append("")  # 空行分隔
        
        return '\n'.join(content)
    
    except ImportError:
        raise ImportError("需要安装pandas和openpyxl库: pip install pandas openpyxl")
    except Exception as e:
        raise Exception(f"读取Excel文档失败: {str(e)}")


def read_file_to_string(file_path: str) -> str:
    """
    根据文件扩展名自动选择合适的读取方法
    
    Args:
        file_path (str): 文件路径
        
    Returns:
        str: 文件内容
        
    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 不支持的文件格式
        Exception: 其他读取错误
    """
    # 检查文件是否存在
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    # 获取文件扩展名
    file_extension = Path(file_path).suffix.lower()
    
    # 根据扩展名选择读取方法
    if file_extension == '.txt':
        return read_txt_file(file_path)
    
    elif file_extension == '.md':
        return read_markdown_file(file_path)
    
    elif file_extension in ['.docx', '.doc']:
        return read_word_file(file_path)
    
    elif file_extension in ['.xlsx', '.xls']:
        return read_excel_file(file_path)
    
    elif file_extension == '.pdf':
        return read_pdf_file(file_path)
    
    else:
        raise ValueError(f"不支持的文件格式: {file_extension}")


def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    获取文件基本信息
    
    Args:
        file_path (str): 文件路径
        
    Returns:
        dict: 文件信息
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    file_stat = os.stat(file_path)
    file_path_obj = Path(file_path)
    
    return {
        'filename': file_path_obj.name,
        'extension': file_path_obj.suffix,
        'size_bytes': file_stat.st_size,
        'size_mb': round(file_stat.st_size / (1024 * 1024), 2),
        'absolute_path': file_path_obj.absolute(),
        'is_supported': file_path_obj.suffix.lower() in ['.txt', '.md', '.docx', '.doc', '.xlsx', '.xls', '.pdf']
    }


# 示例使用函数
def read_pdf_file(file_path: str) -> str:
    """
    读取PDF文档
    
    Args:
        file_path (str): 文件路径
        
    Returns:
        str: 文件内容
        
    Raises:
        ImportError: 需要安装PyPDF2库
        Exception: 读取PDF失败
    """
    try:
        from PyPDF2 import PdfReader
        
        content = []
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                content.append(page.extract_text())
        
        return '\n'.join(content)
    
    except ImportError:
        raise ImportError("需要安装PyPDF2库: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"读取PDF文档失败: {str(e)}")


def demo():
    """
    演示模块功能
    """
    print("=== readfile_2_str.py 演示 ===")
    
    # 示例文件路径（需要根据实际情况修改）
    test_files = [
        "sample_doc/11、《中华人民共和国电信条例》.txt",
        "sample_doc/简化方案说明.md", 
        "sample_doc/11、《中华人民共和国电信条例》.docx",
        "sample_doc/27、用户申诉责任认定标准（V2.0）（正式印发）-20250327.xlsx",
        "sample_doc/云趣运维文档1754623245145/语音机器人安装手册V1.7.pdf"
    ]
    
    for file_path in test_files:
        try:
            print(f"\n--- 读取文件: {file_path} ---")
            
            # 获取文件信息
            info = get_file_info(file_path)
            print(f"文件大小: {info['size_mb']} MB")
            print(f"是否支持: {info['is_supported']}")
            
            # 读取文件内容
            content = read_file_to_string(file_path)
            print(f"内容长度: {len(content)} 字符")
            print(f"内容预览: {content[:100]}...")
            file_name=file_path+".txt"
            with open(file_name,'w',encoding='utf-8') as f:
                f.write(content)

        except Exception as e:
            print(f"错误: {str(e)}")


if __name__ == "__main__":
    demo()
